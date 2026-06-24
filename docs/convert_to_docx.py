"""Convert LAPORAN_TUBES_PBO_SPACEBOOK.md to Word (.docx) with academic formatting."""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn

MD_FILE = Path(__file__).parent / "LAPORAN_TUBES_PBO_SPACEBOOK.md"
OUT_FILE = Path(__file__).parent / "LAPORAN_TUBES_PBO_SPACEBOOK.docx"


def set_document_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12)


def add_cover_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAPORAN TUGAS BESAR\nPEMROGRAMAN BERORIENTASI OBJEK")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Times New Roman"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SpaceBook\nSistem Reservasi Co-Working Space Kampus")
    run.bold = True
    run.font.size = Pt(14)

    for text in [
        "",
        "Kelompok: Penerbang Roket",
        "Program Studi Informatika",
        "Telkom University Purwokerto",
        "Semester Genap 2025/2026",
        "",
        "Dosen Pengampu:",
        "Dany Candra Febrianto, S.Kom., M.Kom.",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    doc.add_page_break()


def add_table_from_md(doc: Document, lines: list[str], start: int) -> int:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= {"-", ":"} for c in row):
            rows.append(row)
        i += 1

    if not rows:
        return start

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            table.rows[ri].cells[ci].text = cell
            for paragraph in table.rows[ri].cells[ci].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
    doc.add_paragraph()
    return i


def add_code_block(doc: Document, lines: list[str], start: int) -> int:
    i = start + 1
    code_lines = []
    while i < len(lines) and not lines[i].strip().startswith("```"):
        code_lines.append(lines[i])
        i += 1
    p = doc.add_paragraph()
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    doc.add_paragraph()
    return i + 1


def add_paragraph_text(doc: Document, text: str, bold: bool = False, center: bool = False) -> None:
    if not text.strip():
        doc.add_paragraph()
        return
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold


def convert_md_to_docx() -> None:
    content = MD_FILE.read_text(encoding="utf-8")
    lines = content.splitlines()

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    set_document_styles(doc)
    add_cover_page(doc)

    skip_until_bab = True
    in_code = False
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if skip_until_bab:
            if stripped.startswith("# BAB I"):
                skip_until_bab = False
            else:
                i += 1
                continue

        if stripped.startswith("```"):
            if not in_code:
                i = add_code_block(doc, lines, i)
                continue
            in_code = False
            i += 1
            continue

        if stripped.startswith("<div") or stripped.startswith("---") and "page-break" in stripped:
            if "page-break" in stripped:
                doc.add_page_break()
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("|"):
            i = add_table_from_md(doc, lines, i)
            continue

        if stripped.startswith("# BAB"):
            title = stripped.lstrip("# ").strip()
            doc.add_page_break()
            doc.add_heading(title, level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            add_paragraph_text(doc, stripped.strip("*"), bold=True)
            i += 1
            continue

        if stripped.startswith("*—") or stripped.startswith("*"):
            add_paragraph_text(doc, stripped.lstrip("*").strip())
            i += 1
            continue

        if stripped:
            # Inline bold handling simple
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.75)
            p.paragraph_format.line_spacing = 1.5
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
        else:
            doc.add_paragraph()

        i += 1

    doc.save(OUT_FILE)
    print(f"Berhasil: {OUT_FILE}")


if __name__ == "__main__":
    convert_md_to_docx()
